"""Validacoes de curriculo gerado em DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from radar_vagas.models import ResumeArtifact
from radar_vagas.resumes.profile import ResumeProfile
from radar_vagas.text_utils import normalize_text

REQUIRED_SECTION_OPTIONS = [
    ["Resumo Profissional"],
    ["Competencias Tecnicas"],
    ["Experiencia Profissional", "Experiencia em Dados"],
    ["Projetos Selecionados"],
    ["Formacao e Destaques"],
]
PLACEHOLDER_MARKERS = ["todo", "placeholder", "<empresa>", "<cargo>"]


def validate_resume_artifact(
    artifact: ResumeArtifact,
    resume_profile: ResumeProfile,
) -> list[str]:
    """Valida o arquivo gerado, seu conteudo e suas restricoes de ATS."""
    file_path = Path(artifact.file_path)
    errors: list[str] = []

    if file_path.suffix.lower() != ".docx":
        errors.append("resume file must use .docx extension")
        return errors
    if not file_path.exists():
        errors.append("resume file was not created")
        return errors

    try:
        document = Document(file_path)
    except Exception:
        errors.append("resume file cannot be opened by python-docx")
        return errors

    document_text = "\n".join(
        paragraph.text for paragraph in document.paragraphs if paragraph.text
    ).strip()
    normalized_text = normalize_text(document_text)

    for section_options in REQUIRED_SECTION_OPTIONS:
        normalized_options = [normalize_text(section) for section in section_options]
        if not any(option in normalized_text for option in normalized_options):
            errors.append(f"missing required section: {section_options[0]}")

    if not document_text:
        errors.append("resume document is empty")

    for marker in PLACEHOLDER_MARKERS:
        if marker in normalized_text:
            errors.append(f"placeholder detected: {marker}")

    forbidden_claims = [
        normalize_text(claim)
        for claim in resume_profile.candidate_profile.forbidden_claims
    ]
    for forbidden_claim in forbidden_claims:
        if forbidden_claim and forbidden_claim in normalized_text:
            errors.append(f"forbidden claim detected: {forbidden_claim}")

    if document.tables:
        errors.append("resume must not contain tables")
    if document.inline_shapes:
        errors.append("resume must not contain images or shapes")

    mentions_esic = normalize_text("ESIC") in normalized_text
    mentions_administracao = normalize_text("Administracao") in normalized_text
    mentions_interrupted = normalize_text("Curso interrompido") in normalized_text
    if (mentions_esic or mentions_administracao) and not mentions_interrupted:
        errors.append("ESIC or Administracao must remain marked as Curso interrompido")

    return errors
