"""Geracao deterministica de curriculos ATS em DOCX."""

from __future__ import annotations

from datetime import datetime
from hashlib import sha256
from pathlib import Path
from urllib.parse import urlsplit

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt

from radar_vagas.config import ProfileConfig
from radar_vagas.models import (
    CandidateAdditionalExperience,
    CandidateEducation,
    CandidateSkill,
    JobPosting,
    ResumeArtifact,
    utc_now,
)
from radar_vagas.resumes.content_selector import SelectedResumeContent
from radar_vagas.resumes.file_names import build_resume_file_name
from radar_vagas.resumes.keyword_extractor import extract_job_keywords
from radar_vagas.resumes.profile import ResumeProfile
from radar_vagas.resumes.validator import validate_resume_artifact
from radar_vagas.text_utils import make_job_fingerprint, normalize_text


def generate_resume(
    *,
    job: JobPosting,
    content: SelectedResumeContent,
    resume_profile: ResumeProfile,
    config: ProfileConfig,
    output_directory: Path | None = None,
) -> ResumeArtifact:
    """Gera e valida um curriculo DOCX para uma vaga elegivel."""
    directory = output_directory or config.resume.output_directory
    target_directory = Path(directory)
    target_directory.mkdir(parents=True, exist_ok=True)

    file_name = build_resume_file_name(
        company=job.company or "Empresa",
        title=content.target_title,
        pattern=config.resume.file_name_pattern,
    )
    file_path = target_directory / file_name

    document = Document()
    _configure_document_styles(document)
    _build_document(document, content, resume_profile)
    file_path = _save_document_with_fallback(document, file_path)
    file_name = file_path.name

    file_hash = sha256(file_path.read_bytes()).hexdigest()
    artifact = ResumeArtifact(
        job_fingerprint=make_job_fingerprint(job.title, job.company, job.location),
        target_title=content.target_title,
        company=job.company or "Empresa",
        file_path=file_path,
        file_name=file_name,
        file_sha256=file_hash,
        selected_skill_ids=content.selected_skill_ids,
        selected_experience_bullet_ids=[
            bullet_id
            for experience in content.experiences
            for bullet_id in experience.bullet_ids
        ],
        selected_project_ids=[project.project.id for project in content.projects],
        generated_at=utc_now(),
        validation_errors=[],
        is_valid=False,
    )
    validation_errors = validate_resume_artifact(artifact, resume_profile)
    artifact.validation_errors = validation_errors
    artifact.is_valid = not validation_errors
    return artifact


def generate_resume_for_job(
    *,
    job: JobPosting,
    resume_profile: ResumeProfile,
    config: ProfileConfig,
    output_directory: Path | None = None,
) -> ResumeArtifact:
    """Executa extracao, selecao, geracao e validacao em uma unica chamada."""
    from radar_vagas.resumes.content_selector import select_resume_content

    extraction = extract_job_keywords(job, config)
    content = select_resume_content(
        job=job,
        extraction=extraction,
        resume_profile=resume_profile,
        config=config,
    )
    return generate_resume(
        job=job,
        content=content,
        resume_profile=resume_profile,
        config=config,
        output_directory=output_directory,
    )


def _configure_document_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.48)
    section.right_margin = Inches(0.48)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(9.5)
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.line_spacing = 1.0

    if "Resume Section" not in document.styles:
        section_style = document.styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = "Aptos"
        section_style.font.size = Pt(10)
        section_style.font.bold = True
        section_style.paragraph_format.space_before = Pt(5)
        section_style.paragraph_format.space_after = Pt(2)

    if "Job Header" not in document.styles:
        job_style = document.styles.add_style("Job Header", WD_STYLE_TYPE.PARAGRAPH)
        job_style.font.name = "Aptos"
        job_style.font.size = Pt(9.5)
        job_style.font.bold = True
        job_style.paragraph_format.space_before = Pt(2)
        job_style.paragraph_format.space_after = Pt(0)

    if "Resume Bullet" not in document.styles:
        bullet_style = document.styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.name = "Aptos"
        bullet_style.font.size = Pt(9)
        bullet_style.paragraph_format.left_indent = Pt(11.5)
        bullet_style.paragraph_format.first_line_indent = Pt(-8.65)
        bullet_style.paragraph_format.space_after = Pt(0.6)


def _build_document(
    document: Document,
    content: SelectedResumeContent,
    resume_profile: ResumeProfile,
) -> None:
    _add_header(document, content, resume_profile)
    _add_section(document, "RESUMO PROFISSIONAL")
    summary = document.add_paragraph(content.summary)
    summary.paragraph_format.space_after = Pt(1)

    _add_section(document, "COMPETENCIAS TECNICAS")
    _add_grouped_skills(document, content, resume_profile)

    _add_section(document, "EXPERIENCIA EM DADOS")
    for selected_experience in content.experiences:
        header = document.add_paragraph(style="Job Header")
        header.add_run(
            f"{selected_experience.experience.company}  |  "
            f"{selected_experience.experience.role}"
        )
        date_range = _format_date_range(
            selected_experience.experience.start_date,
            selected_experience.experience.end_date,
        )
        if date_range:
            header.add_run(f"  |  {date_range}").bold = False
        for bullet_text in selected_experience.bullet_texts:
            _add_bullet(document, bullet_text)

    _add_section(document, "PROJETOS SELECIONADOS")
    for selected_project in content.projects:
        paragraph = document.add_paragraph(style="Resume Bullet")
        paragraph.add_run("\u2022 ")
        paragraph.add_run(f"{selected_project.project.title}: ").bold = True
        paragraph.add_run(" ".join(selected_project.bullets))

    _add_section(document, "FORMACAO E DESTAQUES")
    for education_item in _order_education_items(content.education):
        paragraph = document.add_paragraph(style="Job Header")
        paragraph.add_run(f"{education_item.course} - {education_item.institution}")
        education_range = _format_date_range(
            education_item.start_date,
            education_item.end_date,
        )
        if education_range:
            paragraph.add_run(f"  |  {education_range}").bold = False
        if education_item.expected_completion:
            paragraph.add_run(
                f"  |  conclusao prevista: {education_item.expected_completion}"
            ).bold = False
        elif education_item.status:
            paragraph.add_run(f"  |  {education_item.status.lower()}").bold = False
    for highlight in content.highlights:
        _add_bullet(document, highlight, spacing_after=0.4)

    if content.additional_experience:
        _add_section(document, "EXPERIENCIA ADICIONAL")
        for entry in content.additional_experience:
            _add_additional_experience_line(document, entry.experience)


def _add_header(
    document: Document,
    content: SelectedResumeContent,
    resume_profile: ResumeProfile,
) -> None:
    candidate = resume_profile.candidate_profile.candidate
    summary_skills = _header_skill_summary(content)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.add_run(candidate.name).bold = True
    name.runs[0].font.name = "Aptos Display"
    name.runs[0].font.size = Pt(17)
    name.runs[0].text = candidate.name.upper()
    name.paragraph_format.space_after = Pt(1)

    target = document.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target_run = target.add_run(
        f"{content.target_title.upper()}  |  {summary_skills}"
    )
    target_run.bold = True
    target_run.font.size = Pt(10)
    target.paragraph_format.space_after = Pt(2.5)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(
        f"{candidate.city} - {candidate.state}  |  "
        f"{resume_profile.phone}  |  {resume_profile.email}"
    )
    contact.add_run().add_break(WD_BREAK.LINE)
    contact.add_run(
        f"{_display_url(candidate.linkedin_url)}  |  {_display_url(candidate.github_url)}"
    )
    contact.paragraph_format.space_after = Pt(3)


def _add_section(document: Document, title: str) -> None:
    paragraph = document.add_paragraph(title, style="Resume Section")
    paragraph.runs[0].text = title.upper()


def _add_bullet(
    document: Document,
    text: str,
    *,
    spacing_after: float = 0.6,
) -> None:
    paragraph = document.add_paragraph(style="Resume Bullet")
    paragraph.add_run("\u2022 ")
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(spacing_after)


def _header_skill_summary(content: SelectedResumeContent) -> str:
    top_skills = content.skills[:3]
    if not top_skills:
        return "POWER BI  |  SQL  |  PYTHON"
    return "  |  ".join(skill.upper() for skill in top_skills)


def _display_url(value: object) -> str:
    text = str(value)
    parts = urlsplit(text)
    host = parts.netloc.removeprefix("www.")
    return f"{host}{parts.path}".rstrip("/")


def _format_date_range(start_date: str | None, end_date: str | None) -> str:
    start = _format_month_year(start_date)
    end = _format_month_year(end_date) if end_date else "Atual"
    if start and end:
        return f"{start} - {end}"
    return ""


def _format_month_year(value: str | None) -> str:
    if not value:
        return ""
    try:
        parsed = datetime.strptime(value, "%Y-%m")
    except ValueError:
        try:
            parsed = datetime.strptime(value, "%Y-%m-%d")
        except ValueError:
            return value

    month_names = [
        "jan",
        "fev",
        "mar",
        "abr",
        "mai",
        "jun",
        "jul",
        "ago",
        "set",
        "out",
        "nov",
        "dez",
    ]
    return f"{month_names[parsed.month - 1]}/{parsed.year}"


def _add_grouped_skills(
    document: Document,
    content: SelectedResumeContent,
    resume_profile: ResumeProfile,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    skill_map = {
        skill.id: skill for skill in resume_profile.candidate_profile.skills
    }
    selected_ids = content.selected_skill_ids
    segments = _build_skill_segments(selected_ids, skill_map)

    runs_added = 0
    for title, values in segments:
        if runs_added:
            separator = paragraph.add_run("  |  ")
            separator.font.size = Pt(9)
        title_run = paragraph.add_run(f"{title}: ")
        title_run.bold = True
        title_run.font.size = Pt(9)
        values_run = paragraph.add_run(values)
        values_run.font.size = Pt(9)
        runs_added += 1


def _build_skill_segments(
    selected_ids: list[str],
    skill_map: dict[str, CandidateSkill],
) -> list[tuple[str, str]]:
    normalized_ids = set(selected_ids)
    segments: list[tuple[str, str]] = []

    business_intelligence = _build_business_intelligence_segment(
        skill_map,
        normalized_ids,
    )
    if business_intelligence:
        segments.append(("Business Intelligence", business_intelligence))

    dados_items = _build_data_segment(skill_map, normalized_ids)
    if dados_items:
        segments.append(("Dados", dados_items))

    programacao = _build_programming_segment(skill_map, normalized_ids)
    if programacao:
        segments.append(("Programacao", programacao))

    ferramentas = _build_tools_segment(skill_map)
    if ferramentas:
        segments.append(("Ferramentas", ferramentas))

    if not segments:
        fallback = [
            skill_map[skill_id].label
            for skill_id in selected_ids
            if skill_id in skill_map
        ]
        if fallback:
            segments.append(("Competencias", ", ".join(fallback[:6])))

    return segments


def _skill_label(
    skill_map: dict[str, CandidateSkill],
    skill_id: str,
    selected_ids: set[str],
) -> str | None:
    if skill_id not in selected_ids:
        return None
    skill = skill_map.get(skill_id)
    return None if skill is None else str(skill.label)


def _extract_aliases(
    skill_map: dict[str, CandidateSkill],
    skill_id: str,
    selected_ids: set[str],
    alias_whitelist: set[str],
    *,
    tag_whitelist: set[str] | None = None,
) -> str | None:
    if skill_id not in selected_ids:
        return None
    skill = skill_map.get(skill_id)
    if skill is None:
        return None

    values: list[str] = []
    for alias in skill.aliases:
        if normalize_text(alias) in alias_whitelist:
            values.append(_title_case_skill(alias))
    for tag in skill.tags:
        if tag_whitelist and normalize_text(tag) in tag_whitelist:
            values.append(_title_case_skill(tag))

    unique_values: list[str] = []
    for value in values:
        if value not in unique_values:
            unique_values.append(value)
    return ", ".join(unique_values) or None


def _build_business_intelligence_segment(
    skill_map: dict[str, CandidateSkill],
    selected_ids: set[str],
) -> str | None:
    if not any(skill_id in selected_ids for skill_id in {"power_bi", "dax", "power_query"}):
        return None
    return _unique_join(
        _skill_label(skill_map, "power_bi", selected_ids),
        _skill_label(skill_map, "dax", selected_ids),
        _skill_label(skill_map, "power_query", selected_ids),
        "modelagem de dados e dashboards",
    )


def _build_data_segment(
    skill_map: dict[str, CandidateSkill],
    selected_ids: set[str],
) -> str | None:
    if "sql" not in selected_ids:
        return None
    return _unique_join(
        _skill_label(skill_map, "sql", selected_ids),
        _extract_aliases(skill_map, "sql", selected_ids, {"postgresql", "firebird"}),
        "tratamento, validacao e qualidade de dados",
    )


def _build_programming_segment(
    skill_map: dict[str, CandidateSkill],
    selected_ids: set[str],
) -> str | None:
    if not any(skill_id in selected_ids for skill_id in {"python", "etl"}):
        return None
    python_value = _skill_label(skill_map, "python", selected_ids) or "Python"
    python_details = _build_python_details(skill_map, selected_ids)
    return _append_parenthetical(python_value, python_details)


def _build_tools_segment(skill_map: dict[str, CandidateSkill]) -> str | None:
    if "git" not in skill_map or "excel" not in skill_map:
        return None
    return "Git/GitHub e Excel"


def _build_python_details(
    skill_map: dict[str, CandidateSkill],
    selected_ids: set[str],
) -> str | None:
    if not any(skill_id in selected_ids for skill_id in {"python", "etl"}):
        return None

    python_skill = skill_map.get("python")
    if python_skill is None:
        return "automacoes e ETL"

    detail_parts: list[str] = []
    alias_values = {normalize_text(alias) for alias in python_skill.aliases}
    tag_values = {normalize_text(tag) for tag in python_skill.tags}

    if "pandas" in alias_values:
        detail_parts.append("Pandas")

    has_automation = "automacao" in tag_values
    has_etl = "etl" in tag_values or "etl" in selected_ids
    if has_automation and has_etl:
        detail_parts.append("automacoes e ETL")
    elif has_automation:
        detail_parts.append("automacoes")
    elif has_etl:
        detail_parts.append("ETL")

    return ", ".join(detail_parts) or None


def _append_parenthetical(value: str | None, details: str | None) -> str | None:
    if not value:
        return None
    if not details:
        return value
    return f"{value} ({details})"


def _title_case_skill(value: str) -> str:
    raw = normalize_text(value)
    overrides = {
        "etl": "ETL",
        "sql": "SQL",
        "postgresql": "PostgreSQL",
        "power bi": "Power BI",
        "power query": "Power Query",
        "git": "Git",
        "github": "GitHub",
    }
    if raw in overrides:
        return overrides[raw]
    return " ".join(part.capitalize() for part in value.split())


def _unique_join(*values: str | None) -> str | None:
    collected: list[str] = []
    for value in values:
        if not value or value in collected:
            continue
        collected.append(value)
    return ", ".join(collected) or None


def _order_education_items(
    education_items: list[CandidateEducation],
) -> list[CandidateEducation]:
    def ranking(item: CandidateEducation) -> tuple[int, str]:
        status = normalize_text(item.status)
        if "interrompido" in status:
            return (0, item.course)
        if "andamento" in status:
            return (1, item.course)
        return (2, item.course)

    return sorted(education_items, key=ranking)


def _add_additional_experience_line(
    document: Document,
    experience: CandidateAdditionalExperience,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    headline_run = paragraph.add_run(experience.headline)
    headline_run.bold = True
    headline_run.font.size = Pt(9)
    details_run = paragraph.add_run(f"  |  {experience.details}")
    details_run.font.size = Pt(9)


def _save_document_with_fallback(document: Document, target_path: Path) -> Path:
    try:
        document.save(target_path)
        return target_path
    except PermissionError:
        pass

    for index in range(1, 100):
        candidate_path = target_path.with_name(
            f"{target_path.stem}_{index}{target_path.suffix}"
        )
        try:
            document.save(candidate_path)
            return candidate_path
        except PermissionError:
            continue

    raise PermissionError(
        "Could not save resume because every candidate file path is locked under "
        f"{target_path.parent}"
    )
