"""Envio de notificacoes para Discord via webhook."""

from __future__ import annotations

import hashlib
import json
import time
import zipfile
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Callable

import httpx
from docx import Document as DocxDocument

from radar_vagas.config import RuntimeSettings
from radar_vagas.models import (
    EvaluatedJob,
    JobPosting,
    Priority,
    ResumeArtifact,
    WorkMode,
    utc_now,
)

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DISCORD_EMBED_TITLE_LIMIT = 256
DISCORD_EMBED_DESCRIPTION_LIMIT = 4096
DISCORD_EMBED_FIELD_NAME_LIMIT = 256
DISCORD_EMBED_FIELD_VALUE_LIMIT = 1024
MAX_DISCORD_ATTACHMENT_SIZE_BYTES = 8 * 1024 * 1024


class DiscordNotificationError(RuntimeError):
    """Erro-base para falhas de notificacao no Discord."""


class DiscordConfigurationError(DiscordNotificationError):
    """Erro de configuracao para envio ao Discord."""


class DiscordAttachmentError(DiscordNotificationError):
    """Erro de validacao do anexo enviado ao Discord."""


class DiscordRequestError(DiscordNotificationError):
    """Erro ao executar a requisicao HTTP para o Discord."""


@dataclass(slots=True)
class DiscordMessageReceipt:
    """Resumo seguro do resultado de uma mensagem enviada."""

    message_id: str | None
    status_code: int
    attachment_name: str
    attempts: int


def send_job_notification(
    *,
    webhook_url: str,
    evaluated_job: EvaluatedJob,
    resume_artifact: ResumeArtifact,
    timeout_seconds: float = 10.0,
    max_attempts: int = 3,
    client: httpx.Client | None = None,
    sleep_func: Callable[[float], None] = time.sleep,
) -> DiscordMessageReceipt:
    """Envia uma vaga com um curriculo DOCX como anexo para o webhook do Discord."""
    cleaned_webhook_url = webhook_url.strip()
    if not cleaned_webhook_url:
        raise DiscordConfigurationError(
            "DISCORD_WEBHOOK_URL is required to send Discord notifications"
        )

    attachment_path = resume_artifact.file_path
    _validate_docx_attachment(attachment_path)
    attachment_bytes = attachment_path.read_bytes()
    request_url = str(httpx.URL(cleaned_webhook_url).copy_add_param("wait", "true"))
    payload_json = json.dumps(
        _build_webhook_payload(evaluated_job=evaluated_job),
        ensure_ascii=False,
    )

    owns_client = client is None
    http_client = client or httpx.Client(timeout=timeout_seconds)

    try:
        for attempt in range(1, max_attempts + 1):
            try:
                response = http_client.post(
                    request_url,
                    data={"payload_json": payload_json},
                    files={
                        "files[0]": (
                            resume_artifact.file_name,
                            attachment_bytes,
                            DOCX_CONTENT_TYPE,
                        )
                    },
                )
            except httpx.TimeoutException as exc:
                if attempt >= max_attempts:
                    raise DiscordRequestError(
                        "Discord request timed out after the configured retries"
                    ) from exc
                sleep_func(_backoff_seconds(attempt))
                continue
            except httpx.HTTPError as exc:
                if attempt >= max_attempts:
                    raise DiscordRequestError(
                        "Discord request failed due to a transient network error"
                    ) from exc
                sleep_func(_backoff_seconds(attempt))
                continue

            if 200 <= response.status_code < 300:
                message_id = _extract_message_id(response)
                return DiscordMessageReceipt(
                    message_id=message_id,
                    status_code=response.status_code,
                    attachment_name=resume_artifact.file_name,
                    attempts=attempt,
                )

            if response.status_code == 429:
                if attempt >= max_attempts:
                    raise DiscordRequestError("Discord rate limited the request after retries")
                sleep_func(_retry_after_seconds(response, attempt))
                continue

            if response.status_code in {408, 500, 502, 503, 504}:
                if attempt >= max_attempts:
                    raise DiscordRequestError(
                        f"Discord webhook failed with transient status {response.status_code}"
                    )
                sleep_func(_backoff_seconds(attempt))
                continue

            raise DiscordRequestError(
                f"Discord webhook rejected the request with status {response.status_code}"
            )
    finally:
        if owns_client:
            http_client.close()


def send_test_message(
    *,
    settings: RuntimeSettings,
    output_directory: Path | None = None,
    client: httpx.Client | None = None,
    sleep_func: Callable[[float], None] = time.sleep,
) -> DiscordMessageReceipt:
    """Gera um DOCX ficticio e envia uma mensagem de teste para o Discord."""
    if not settings.discord_webhook_url:
        raise DiscordConfigurationError(
            "DISCORD_WEBHOOK_URL is required to test Discord notifications"
        )

    target_directory = output_directory or Path("output/resumes")
    target_directory.mkdir(parents=True, exist_ok=True)
    attachment_path = target_directory / "Curriculo_Teste_Discord_Radar_Vagas.docx"
    _create_test_docx(attachment_path)
    attachment_hash = hashlib.sha256(attachment_path.read_bytes()).hexdigest()

    job = JobPosting(
        provider="diagnostic",
        provider_job_id="discord-test",
        title="Analista de BI Junior",
        company="Empresa Ficticia",
        location="Curitiba - PR",
        work_mode=WorkMode.REMOTE,
        description="Mensagem ficticia para validar o webhook do Discord sem expor dados reais.",
        url="https://example.com/jobs/discord-test",
        source_name="Teste Local",
        updated_at=utc_now(),
    )
    evaluated_job = EvaluatedJob(
        job=job,
        score=92,
        priority=Priority.HIGH,
        matched_skills=["Power BI", "SQL", "Python"],
        missing_skills=["ETL"],
        extracted_keywords=["power bi", "sql", "python"],
        relevant_domains=["bi"],
        rejection_reasons=[],
        is_eligible=True,
        fingerprint="discord-test-fingerprint",
        score_explanation="Mensagem de teste com alta aderencia para validar webhook e anexo.",
    )
    resume_artifact = ResumeArtifact(
        job_fingerprint=evaluated_job.fingerprint,
        target_title=job.title,
        company=job.company or "Empresa Ficticia",
        file_path=attachment_path,
        file_name=attachment_path.name,
        file_sha256=attachment_hash,
        selected_skill_ids=["power-bi", "sql", "python"],
        selected_experience_bullet_ids=["exp-test-1"],
        selected_project_ids=["project-test-1"],
        generated_at=utc_now(),
        validation_errors=[],
        is_valid=True,
    )

    return send_job_notification(
        webhook_url=settings.discord_webhook_url,
        evaluated_job=evaluated_job,
        resume_artifact=resume_artifact,
        client=client,
        sleep_func=sleep_func,
    )


def _build_webhook_payload(*, evaluated_job: EvaluatedJob) -> dict[str, Any]:
    job = evaluated_job.job
    timestamp = job.updated_at or job.published_at or utc_now()
    embed = {
        "title": _truncate(job.title, DISCORD_EMBED_TITLE_LIMIT),
        "url": str(job.url),
        "description": _truncate(
            "Curriculo otimizado para esta vaga anexado nesta mensagem.",
            DISCORD_EMBED_DESCRIPTION_LIMIT,
        ),
        "color": 0x1F8B4C,
        "fields": [
            _build_field("Cargo", job.title),
            _build_field("Empresa", job.company or "Nao informado"),
            _build_field("Localizacao", job.location or "Nao informada"),
            _build_field("Score", f"{evaluated_job.score}/100"),
            _build_field("Prioridade", _format_priority(evaluated_job.priority)),
            _build_field("Competencias aderentes", _join_values(evaluated_job.matched_skills)),
            _build_field("Competencias ausentes", _join_values(evaluated_job.missing_skills)),
            _build_field("Data", _format_datetime(timestamp)),
            _build_field("Fonte", job.source_name),
        ],
        "footer": {"text": _truncate(evaluated_job.score_explanation, 256)},
        "timestamp": timestamp.astimezone(UTC).isoformat(),
    }
    return {
        "allowed_mentions": {"parse": []},
        "embeds": [embed],
    }


def _build_field(name: str, value: str) -> dict[str, str]:
    return {
        "name": _truncate(name, DISCORD_EMBED_FIELD_NAME_LIMIT),
        "value": _truncate(value, DISCORD_EMBED_FIELD_VALUE_LIMIT),
    }


def _join_values(values: list[str]) -> str:
    if not values:
        return "Nenhuma"
    return ", ".join(values)


def _truncate(value: str, limit: int) -> str:
    cleaned = " ".join(value.split())
    if len(cleaned) <= limit:
        return cleaned
    if limit <= 3:
        return cleaned[:limit]
    return f"{cleaned[: limit - 3].rstrip()}..."


def _format_priority(priority: Priority) -> str:
    labels = {
        Priority.HIGH: "Alta",
        Priority.GOOD: "Boa oportunidade",
        Priority.BELOW_THRESHOLD: "Abaixo do minimo",
    }
    return labels[priority]


def _format_datetime(value: datetime) -> str:
    return value.astimezone(UTC).strftime("%d/%m/%Y")


def _extract_message_id(response: httpx.Response) -> str | None:
    try:
        payload = response.json()
    except ValueError:
        return None

    if isinstance(payload, dict) and payload.get("id") is not None:
        return str(payload["id"])
    return None


def _retry_after_seconds(response: httpx.Response, attempt: int) -> float:
    retry_after = _extract_retry_after_from_body(response)
    if retry_after is not None:
        return retry_after

    header_value = response.headers.get("Retry-After")
    if header_value:
        try:
            return max(float(header_value), 0.0)
        except ValueError:
            pass

    return _backoff_seconds(attempt)


def _extract_retry_after_from_body(response: httpx.Response) -> float | None:
    try:
        payload = response.json()
    except ValueError:
        return None

    if not isinstance(payload, dict):
        return None

    raw_value = payload.get("retry_after")
    if raw_value is None:
        return None

    try:
        return max(float(raw_value), 0.0)
    except (TypeError, ValueError):
        return None


def _backoff_seconds(attempt: int) -> float:
    return min(2 ** (attempt - 1), 8)


def _validate_docx_attachment(path: Path) -> None:
    if not path.exists():
        raise DiscordAttachmentError(f"DOCX attachment not found: {path}")

    if path.suffix.lower() != ".docx":
        raise DiscordAttachmentError("Discord attachment must use the .docx extension")

    file_size = path.stat().st_size
    if file_size <= 0:
        raise DiscordAttachmentError("Discord attachment is empty")

    if file_size > MAX_DISCORD_ATTACHMENT_SIZE_BYTES:
        raise DiscordAttachmentError(
            "Discord attachment exceeds the conservative 8 MiB preflight limit"
        )

    try:
        with zipfile.ZipFile(path) as archive:
            if "word/document.xml" not in archive.namelist():
                raise DiscordAttachmentError("DOCX attachment is missing word/document.xml")
    except zipfile.BadZipFile as exc:
        raise DiscordAttachmentError("DOCX attachment is invalid or corrupted") from exc
    except OSError as exc:
        raise DiscordAttachmentError("Could not open the DOCX attachment") from exc


def _create_test_docx(path: Path) -> None:
    document = DocxDocument()
    document.add_heading("Curriculo Ficticio de Teste", level=1)
    document.add_paragraph("Documento gerado para validar o envio de anexos ao Discord.")
    document.add_paragraph("Este arquivo nao contem dados pessoais reais.")
    document.save(path)
