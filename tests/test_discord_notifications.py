from __future__ import annotations

import hashlib
import json
from pathlib import Path

import httpx
import pytest
import respx
from docx import Document as DocxDocument

from radar_vagas.cli import main
from radar_vagas.config import RuntimeSettings
from radar_vagas.models import EvaluatedJob, JobPosting, Priority, ResumeArtifact, WorkMode
from radar_vagas.notifications.discord import (
    DiscordAttachmentError,
    DiscordConfigurationError,
    DiscordMessageReceipt,
    DiscordRequestError,
    send_job_notification,
    send_test_message,
)


def _build_docx(path: Path) -> None:
    document = DocxDocument()
    document.add_heading("Curriculo de Teste", level=1)
    document.add_paragraph("Anexo ficticio para testes do webhook.")
    document.save(path)


def _build_evaluated_job() -> EvaluatedJob:
    job = JobPosting(
        provider="manual",
        provider_job_id="job-123",
        title="Analista de Dados Junior",
        company="Empresa XPTO",
        location="Curitiba - PR",
        work_mode=WorkMode.REMOTE,
        description="SQL, Power BI e Python.",
        url="https://example.com/jobs/123",
        source_name="Teste",
        updated_at="2026-06-29T12:00:00Z",
    )
    return EvaluatedJob(
        job=job,
        score=88,
        priority=Priority.HIGH,
        matched_skills=["Power BI", "SQL", "Python"],
        missing_skills=["ETL"],
        extracted_keywords=["power bi", "sql", "python"],
        relevant_domains=["bi"],
        rejection_reasons=[],
        is_eligible=True,
        fingerprint="fingerprint-123",
        score_explanation="Alta aderencia para teste do envio ao Discord.",
    )


def _build_resume_artifact(path: Path) -> ResumeArtifact:
    file_hash = hashlib.sha256(path.read_bytes()).hexdigest()
    return ResumeArtifact(
        job_fingerprint="fingerprint-123",
        target_title="Analista de Dados Junior",
        company="Empresa XPTO",
        file_path=path,
        file_name=path.name,
        file_sha256=file_hash,
        selected_skill_ids=["power-bi", "sql"],
        selected_experience_bullet_ids=["exp-1"],
        selected_project_ids=["project-1"],
        generated_at="2026-06-29T12:00:00Z",
        validation_errors=[],
        is_valid=True,
    )


@respx.mock
def test_send_job_notification_success_with_attachment(tmp_path: Path) -> None:
    attachment_path = tmp_path / "resume.docx"
    _build_docx(attachment_path)
    evaluated_job = _build_evaluated_job()
    artifact = _build_resume_artifact(attachment_path)
    route = respx.post("https://discord.example/webhook?wait=true").mock(
        return_value=httpx.Response(200, json={"id": "message-1"})
    )

    receipt = send_job_notification(
        webhook_url="https://discord.example/webhook",
        evaluated_job=evaluated_job,
        resume_artifact=artifact,
    )

    assert route.called
    request = route.calls[0].request
    body = request.content.decode("utf-8", errors="ignore")
    assert request.method == "POST"
    assert "multipart/form-data" in request.headers["Content-Type"]
    assert "payload_json" in body
    assert '"parse": []' in body
    assert artifact.file_name in body
    assert receipt.message_id == "message-1"
    assert receipt.status_code == 200


def test_send_job_notification_rejects_missing_attachment(tmp_path: Path) -> None:
    attachment_path = tmp_path / "missing.docx"
    existing_path = tmp_path / "existing.docx"
    _build_docx(existing_path)
    evaluated_job = _build_evaluated_job()
    artifact = _build_resume_artifact(existing_path)
    artifact = artifact.model_copy(
        update={"file_path": attachment_path, "file_name": attachment_path.name}
    )

    with pytest.raises(DiscordAttachmentError, match="not found"):
        send_job_notification(
            webhook_url="https://discord.example/webhook",
            evaluated_job=evaluated_job,
            resume_artifact=artifact,
        )


def test_send_job_notification_rejects_invalid_attachment(tmp_path: Path) -> None:
    attachment_path = tmp_path / "resume.docx"
    attachment_path.write_text("not a docx", encoding="utf-8")
    evaluated_job = _build_evaluated_job()
    artifact = _build_resume_artifact(attachment_path)

    with pytest.raises(DiscordAttachmentError, match="invalid or corrupted"):
        send_job_notification(
            webhook_url="https://discord.example/webhook",
            evaluated_job=evaluated_job,
            resume_artifact=artifact,
        )


@respx.mock
def test_send_job_notification_raises_on_400(tmp_path: Path) -> None:
    attachment_path = tmp_path / "resume.docx"
    _build_docx(attachment_path)
    evaluated_job = _build_evaluated_job()
    artifact = _build_resume_artifact(attachment_path)
    route = respx.post("https://discord.example/webhook?wait=true").mock(
        return_value=httpx.Response(400, json={"message": "bad request"})
    )

    with pytest.raises(DiscordRequestError, match="status 400"):
        send_job_notification(
            webhook_url="https://discord.example/webhook",
            evaluated_job=evaluated_job,
            resume_artifact=artifact,
        )

    assert route.call_count == 1


@respx.mock
def test_send_job_notification_retries_on_429(tmp_path: Path) -> None:
    attachment_path = tmp_path / "resume.docx"
    _build_docx(attachment_path)
    evaluated_job = _build_evaluated_job()
    artifact = _build_resume_artifact(attachment_path)
    slept: list[float] = []
    route = respx.post("https://discord.example/webhook?wait=true").mock(
        side_effect=[
            httpx.Response(429, json={"retry_after": 0.25}),
            httpx.Response(200, json={"id": "message-2"}),
        ]
    )

    receipt = send_job_notification(
        webhook_url="https://discord.example/webhook",
        evaluated_job=evaluated_job,
        resume_artifact=artifact,
        sleep_func=slept.append,
    )

    assert route.call_count == 2
    assert slept == [0.25]
    assert receipt.message_id == "message-2"


@respx.mock
def test_send_job_notification_retries_on_timeout(tmp_path: Path) -> None:
    attachment_path = tmp_path / "resume.docx"
    _build_docx(attachment_path)
    evaluated_job = _build_evaluated_job()
    artifact = _build_resume_artifact(attachment_path)
    slept: list[float] = []
    route = respx.post("https://discord.example/webhook?wait=true").mock(
        side_effect=httpx.TimeoutException("boom")
    )

    with pytest.raises(DiscordRequestError, match="timed out"):
        send_job_notification(
            webhook_url="https://discord.example/webhook",
            evaluated_job=evaluated_job,
            resume_artifact=artifact,
            sleep_func=slept.append,
        )

    assert route.call_count == 3
    assert slept == [1, 2]


def test_send_test_message_requires_discord_configuration(tmp_path: Path) -> None:
    settings = RuntimeSettings.model_validate({"_env_file": None})

    with pytest.raises(DiscordConfigurationError, match="DISCORD_WEBHOOK_URL"):
        send_test_message(settings=settings, output_directory=tmp_path)


@respx.mock
def test_send_test_message_sends_fictitious_docx(tmp_path: Path) -> None:
    settings = RuntimeSettings.model_validate(
        {"discord_webhook_url": "https://discord.example/webhook", "_env_file": None}
    )
    route = respx.post("https://discord.example/webhook?wait=true").mock(
        return_value=httpx.Response(200, json={"id": "message-4"})
    )

    receipt = send_test_message(settings=settings, output_directory=tmp_path)

    assert route.called
    assert (tmp_path / "Curriculo_Teste_Discord_Radar_Vagas.docx").exists()
    assert receipt.message_id == "message-4"


def test_cli_test_discord(
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
) -> None:
    monkeypatch.setattr(
        "radar_vagas.cli.load_runtime_settings",
        lambda: RuntimeSettings.model_validate(
            {"discord_webhook_url": "https://discord.example/webhook", "_env_file": None}
        ),
    )
    monkeypatch.setattr(
        "radar_vagas.cli.send_test_message",
        lambda *, settings: DiscordMessageReceipt(
            message_id="message-3",
            status_code=200,
            attachment_name="Curriculo_Teste_Discord_Radar_Vagas.docx",
            attempts=1,
        ),
    )

    exit_code = main(["--test-discord"])

    captured = capsys.readouterr()
    payload = json.loads(captured.out)
    assert exit_code == 0
    assert payload["message_id"] == "message-3"
    assert payload["status_code"] == 200
