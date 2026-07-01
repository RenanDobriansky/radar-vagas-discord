"""Testes para perfil profissional e geracao deterministica de curriculos."""

from __future__ import annotations

from dataclasses import replace
from pathlib import Path

import pytest
from docx import Document

from radar_vagas.cli import main
from radar_vagas.config import FileConfigurationError, load_candidate_profile, load_profile_config
from radar_vagas.models import JobPosting, ResumeArtifact
from radar_vagas.resumes.content_selector import (
    SelectedExperience,
    SelectedProject,
    select_resume_content,
)
from radar_vagas.resumes.file_names import build_resume_file_name
from radar_vagas.resumes.generator import generate_resume, generate_resume_for_job
from radar_vagas.resumes.keyword_extractor import extract_job_keywords
from radar_vagas.resumes.profile import ResumeProfile, load_resume_profile
from radar_vagas.resumes.validator import validate_resume_artifact
from radar_vagas.text_utils import normalize_text

FIXTURES_DIR = Path(__file__).parent / "fixtures"
JOBS_DIR = FIXTURES_DIR / "jobs"
CANDIDATE_PROFILE_PATH = FIXTURES_DIR / "candidate_profile.yaml"


@pytest.fixture
def profile_config():
    return load_profile_config()


@pytest.fixture
def resume_profile() -> ResumeProfile:
    class SettingsStub:
        candidate_email = "renan.ficticio@example.com"
        candidate_phone = "+55 41 99999-9999"
        candidate_profile_path = CANDIDATE_PROFILE_PATH

    return load_resume_profile(SettingsStub())


def load_job_fixture(file_name: str) -> JobPosting:
    path = JOBS_DIR / file_name
    return JobPosting.model_validate_json(path.read_text(encoding="utf-8"))


def load_document_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)


def build_minimal_resume_artifact(file_path: Path) -> ResumeArtifact:
    return ResumeArtifact(
        job_fingerprint="abc123",
        target_title="Analista de Dados",
        company="Empresa Teste",
        file_path=file_path,
        file_name=file_path.name,
        file_sha256="deadbeef",
        is_valid=False,
    )


def test_load_resume_profile_requires_contact_info() -> None:
    class SettingsStub:
        candidate_email = ""
        candidate_phone = ""
        candidate_profile_path = CANDIDATE_PROFILE_PATH

    with pytest.raises(
        FileConfigurationError,
        match="CANDIDATE_EMAIL is required to generate resumes",
    ):
        load_resume_profile(SettingsStub())


def test_select_resume_content_orders_skills_for_bi_job(
    profile_config,
    resume_profile: ResumeProfile,
) -> None:
    job = load_job_fixture("bi_job.json")
    extraction = extract_job_keywords(job, profile_config)

    content = select_resume_content(
        job=job,
        extraction=extraction,
        resume_profile=resume_profile,
        config=profile_config,
    )

    assert content.target_title == "Analista de BI Junior"
    assert content.selected_skill_ids[0] == "power_bi"
    assert {"dax", "power_query", "sql"} <= set(content.selected_skill_ids[:4])
    assert "Cloud" not in content.skills
    assert "Power BI" in content.skills


def test_select_resume_content_picks_relevant_bullets(
    profile_config,
    resume_profile: ResumeProfile,
) -> None:
    job = load_job_fixture("bi_job.json")
    extraction = extract_job_keywords(job, profile_config)

    content = select_resume_content(
        job=job,
        extraction=extraction,
        resume_profile=resume_profile,
        config=profile_config,
    )

    bullet_ids = [
        bullet_id
        for experience in content.experiences
        for bullet_id in experience.bullet_ids
    ]

    assert "smart_dashboard_indicators" in bullet_ids
    assert "smart_dax_power_query" in bullet_ids
    assert len(bullet_ids) <= 3


@pytest.mark.parametrize(
    ("job_file", "expected_summary_snippet", "expected_project_id"),
    [
        (
            "finance_job.json",
            "indicadores financeiros",
            "cvm_finance_pipeline",
        ),
        (
            "bi_job.json",
            "experiencia em BI, SQL, Power BI",
            "bi_erp_firebird",
        ),
        (
            "data_engineering_job.json",
            "integracao de dados, ETL e automacoes",
            "cvm_finance_pipeline",
        ),
    ],
)
def test_select_resume_content_for_target_domains(
    profile_config,
    resume_profile: ResumeProfile,
    job_file: str,
    expected_summary_snippet: str,
    expected_project_id: str,
) -> None:
    job = load_job_fixture(job_file)
    extraction = extract_job_keywords(job, profile_config)

    content = select_resume_content(
        job=job,
        extraction=extraction,
        resume_profile=resume_profile,
        config=profile_config,
    )

    assert expected_summary_snippet in content.summary
    assert 2 <= len(content.projects) <= 3
    assert expected_project_id in [project.project.id for project in content.projects]


def test_build_resume_file_name_sanitizes_special_characters() -> None:
    file_name = build_resume_file_name(
        company='Empresa: BI/Financeiro?*',
        title='Analista <> Dados | BI',
        pattern="Curriculo_Renan_Dobriansky_{company}_{title}.docx",
    )

    assert file_name == "Curriculo_Renan_Dobriansky_Empresa_BI_Financeiro_Analista_Dados_BI.docx"


def test_generate_resume_can_be_opened_again(
    tmp_path: Path,
    profile_config,
    resume_profile: ResumeProfile,
) -> None:
    job = load_job_fixture("bi_job.json")

    artifact = generate_resume_for_job(
        job=job,
        resume_profile=resume_profile,
        config=profile_config,
        output_directory=tmp_path,
    )

    text = load_document_text(artifact.file_path)
    normalized_text = normalize_text(text)

    assert artifact.is_valid is True
    assert artifact.validation_errors == []
    assert "resumo profissional" in normalized_text
    assert "competencias tecnicas" in normalized_text
    assert "curso interrompido" in normalized_text
    assert "administracao" in normalized_text
    assert "experiencia adicional" in normalized_text
    assert "abr/2025 - dez/2025" in normalized_text
    assert "cursos de power bi e python" in normalized_text


def test_generate_resume_includes_profile_based_skill_groups(
    tmp_path: Path,
    profile_config,
    resume_profile: ResumeProfile,
) -> None:
    job = load_job_fixture("data_engineering_job.json")

    artifact = generate_resume_for_job(
        job=job,
        resume_profile=resume_profile,
        config=profile_config,
        output_directory=tmp_path,
    )

    text = load_document_text(artifact.file_path)
    normalized_text = normalize_text(text)

    assert (
        "dados: sql, postgresql, firebird, sql server, tratamento, "
        "validacao, qualidade de dados"
    ) in normalized_text
    assert "programacao: python (pandas, automacao, etl)" in normalized_text
    assert "modelagem de dados e dashboards" not in normalized_text
    assert "git/github e excel" not in normalized_text


def test_generate_resume_prioritizes_current_data_science_education(
    tmp_path: Path,
    profile_config,
    resume_profile: ResumeProfile,
) -> None:
    job = load_job_fixture("bi_job.json")

    artifact = generate_resume_for_job(
        job=job,
        resume_profile=resume_profile,
        config=profile_config,
        output_directory=tmp_path,
    )

    normalized_text = normalize_text(load_document_text(artifact.file_path))

    assert normalized_text.index("ciencia de dados para negocios") < normalized_text.index(
        "administracao"
    )


def test_generate_resume_can_omit_esic_when_budget_is_tight(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    profile_config,
    resume_profile: ResumeProfile,
) -> None:
    job = load_job_fixture("bi_job.json")
    extraction = extract_job_keywords(job, profile_config)
    content = select_resume_content(
        job=job,
        extraction=extraction,
        resume_profile=resume_profile,
        config=profile_config,
    )
    crowded_content = replace(
        content,
        experiences=[
            SelectedExperience(
                experience=content.experiences[0].experience,
                bullet_ids=[
                    *content.experiences[0].bullet_ids,
                    "extra_bullet_1",
                    "extra_bullet_2",
                ],
                bullet_texts=[
                    *content.experiences[0].bullet_texts,
                    "Projeto adicional validado no perfil para ampliar a aderencia.",
                    "Atividade adicional validada no perfil para ampliar a aderencia.",
                ],
            )
        ],
        projects=[
            *content.projects,
            SelectedProject(
                project=content.projects[0].project,
                bullets=[
                    *content.projects[0].bullets,
                    "Bullet adicional aprovado para pressionar o orcamento.",
                ],
            ),
        ],
        highlights=[
            *content.highlights,
            "Destaque adicional aprovado para pressionar o orcamento.",
        ],
    )
    monkeypatch.setattr("radar_vagas.resumes.generator.ESTIMATED_LINES_PER_PAGE", 12)

    artifact = generate_resume(
        job=job,
        content=crowded_content,
        resume_profile=resume_profile,
        config=profile_config,
        output_directory=tmp_path,
    )

    normalized_text = normalize_text(load_document_text(artifact.file_path))

    assert artifact.is_valid is True
    assert "ciencia de dados para negocios" in normalized_text
    assert "administracao" not in normalized_text
    assert "esic" not in normalized_text
    assert "curso interrompido" not in normalized_text


def test_generate_resume_budget_removes_secondary_content_first(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    profile_config,
    resume_profile: ResumeProfile,
) -> None:
    job = load_job_fixture("data_engineering_job.json")
    extraction = extract_job_keywords(job, profile_config)
    content = select_resume_content(
        job=job,
        extraction=extraction,
        resume_profile=resume_profile,
        config=profile_config,
    )
    monkeypatch.setattr("radar_vagas.resumes.generator.ESTIMATED_LINES_PER_PAGE", 18)

    artifact = generate_resume(
        job=job,
        content=content,
        resume_profile=resume_profile,
        config=profile_config,
        output_directory=tmp_path,
    )

    normalized_text = normalize_text(load_document_text(artifact.file_path))

    assert "experiencia adicional" not in normalized_text
    assert "ciencia de dados para negocios" in normalized_text
    assert "empresa analitica exemplo" in normalized_text


def test_generate_resume_budget_reduces_projects_and_secondary_bullets(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    profile_config,
    resume_profile: ResumeProfile,
) -> None:
    job = load_job_fixture("data_engineering_job.json")
    extraction = extract_job_keywords(job, profile_config)
    content = select_resume_content(
        job=job,
        extraction=extraction,
        resume_profile=resume_profile,
        config=profile_config,
    )
    expanded_content = replace(
        content,
        experiences=[
            SelectedExperience(
                experience=content.experiences[0].experience,
                bullet_ids=[
                    *content.experiences[0].bullet_ids,
                    "extra_bullet_1",
                    "extra_bullet_2",
                ],
                bullet_texts=[
                    *content.experiences[0].bullet_texts,
                    "Experiencia adicional aprovada para ampliar o teste de orcamento.",
                    "Mais uma experiencia aprovada para ampliar o teste de orcamento.",
                ],
            )
        ],
        projects=[
            SelectedProject(
                project=project.project,
                bullets=[
                    *project.bullets,
                    "Bullet adicional aprovado para ampliar o teste de orcamento.",
                ],
            )
            for project in content.projects
        ],
    )
    monkeypatch.setattr("radar_vagas.resumes.generator.ESTIMATED_LINES_PER_PAGE", 14)

    artifact = generate_resume(
        job=job,
        content=expanded_content,
        resume_profile=resume_profile,
        config=profile_config,
        output_directory=tmp_path,
    )

    assert len(artifact.selected_project_ids) == 2
    assert len(artifact.selected_experience_bullet_ids) < len(
        expanded_content.experiences[0].bullet_ids
    )


def test_validate_resume_detects_corrupted_file(
    tmp_path: Path,
    resume_profile: ResumeProfile,
) -> None:
    file_path = tmp_path / "corrompido.docx"
    file_path.write_bytes(b"nao-e-um-docx-valido")

    artifact = build_minimal_resume_artifact(file_path)

    errors = validate_resume_artifact(artifact, resume_profile)

    assert "resume file cannot be opened by python-docx" in errors


def test_validate_resume_requires_interrupted_marker_only_when_esic_or_administracao_appear(
    tmp_path: Path,
    resume_profile: ResumeProfile,
) -> None:
    file_path = tmp_path / "esic_sem_interrupcao.docx"
    document = Document()
    for section in [
        "Resumo Profissional",
        "Competencias Tecnicas",
        "Experiencia em Dados",
        "Projetos Selecionados",
        "Formacao e Destaques",
    ]:
        document.add_paragraph(section)
    document.add_paragraph("Administracao - ESIC Exemplo de Gestao")
    document.save(file_path)

    errors = validate_resume_artifact(build_minimal_resume_artifact(file_path), resume_profile)

    assert "ESIC or Administracao must remain marked as Curso interrompido" in errors


def test_validate_resume_accepts_document_without_esic(
    tmp_path: Path,
    resume_profile: ResumeProfile,
) -> None:
    file_path = tmp_path / "sem_esic.docx"
    document = Document()
    for section in [
        "Resumo Profissional",
        "Competencias Tecnicas",
        "Experiencia em Dados",
        "Projetos Selecionados",
        "Formacao e Destaques",
    ]:
        document.add_paragraph(section)
    document.add_paragraph("Ciencia de Dados para Negocios - Faculdade Exemplo de Negocios")
    document.save(file_path)

    errors = validate_resume_artifact(build_minimal_resume_artifact(file_path), resume_profile)

    assert "ESIC or Administracao must remain marked as Curso interrompido" not in errors


def test_validate_resume_detects_forbidden_claim(
    tmp_path: Path,
    resume_profile: ResumeProfile,
) -> None:
    file_path = tmp_path / "claim_proibido.docx"
    document = Document()
    for section in [
        "Resumo Profissional",
        "Competencias Tecnicas",
        "Experiencia em Dados",
        "Projetos Selecionados",
        "Formacao e Destaques",
    ]:
        document.add_paragraph(section)
    document.add_paragraph("Formado em Administracao com experiencia analitica.")
    document.save(file_path)

    errors = validate_resume_artifact(build_minimal_resume_artifact(file_path), resume_profile)

    assert "forbidden claim detected: formado em administracao" in errors


def test_invalid_candidate_profile_raises_error(tmp_path: Path) -> None:
    invalid_profile = tmp_path / "candidate_profile_invalid.yaml"
    invalid_profile.write_text(
        "\n".join(
            [
                "candidate:",
                '  name: "Pessoa Ficticia"',
                '  city: "Curitiba"',
                '  state: "PR"',
                '  linkedin_url: "https://example.com/linkedin"',
                '  github_url: "https://example.com/github"',
                "summary_blocks:",
                '  - id: "duplicado"',
                '    text: "Resumo 1"',
                '    tags: ["dados"]',
                "skills:",
                '  - id: "duplicado"',
                '    label: "SQL"',
                '    aliases: ["sql"]',
                '    tags: ["sql"]',
                "experiences:",
                '  - id: "exp_1"',
                '    company: "Empresa X"',
                '    role: "Analista"',
                "    bullets:",
                '      - id: "bullet_1"',
                '        text: "Atuacao com dados."',
                '        tags: ["dados"]',
                "projects:",
                '  - id: "proj_1"',
                '    title: "Projeto X"',
                '    bullets: ["Bullet valido"]',
                '    tags: ["dados"]',
                "education:",
                '  - id: "edu_1"',
                '    institution: "ESIC"',
                '    course: "Administracao"',
                '    status: "Concluido"',
                "highlights:",
                '  - "Destaque"',
                "forbidden_claims:",
                '  - "algo proibido"',
            ]
        ),
        encoding="utf-8",
    )

    with pytest.raises(FileConfigurationError, match="Invalid candidate profile"):
        load_candidate_profile(invalid_profile)


def test_cli_generate_resume_command(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("CANDIDATE_EMAIL", "renan.ficticio@example.com")
    monkeypatch.setenv("CANDIDATE_PHONE", "+55 41 99999-9999")
    monkeypatch.setenv("CANDIDATE_PROFILE_PATH", str(CANDIDATE_PROFILE_PATH))

    result = main(
        [
            "--generate-resume",
            str(JOBS_DIR / "bi_job.json"),
        ]
    )

    assert result == 0


def test_cli_dry_run_save_resumes(
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
) -> None:
    from radar_vagas.pipeline import PipelineSummary

    monkeypatch.setattr(
        "radar_vagas.cli.load_runtime_settings",
        lambda: type(
            "SettingsStub",
            (),
            {
                "discord_webhook_url": None,
            },
        )(),
    )
    monkeypatch.setattr(
        "radar_vagas.cli.run_pipeline",
        lambda *, options, settings: PipelineSummary(
            dry_run=options.dry_run,
            provider_names=options.provider_names or ["jooble", "remotive"],
            selected_jobs=0,
        ),
    )

    result = main(["--dry-run", "--save-resumes"])
    captured = capsys.readouterr()

    assert result == 0
    assert '"dry_run": true' in captured.out.lower()
